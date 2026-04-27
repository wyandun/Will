from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════
# HELPERS — same as project_charter_generator.py
# ═══════════════════════════════════════════════

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='CCCCCC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)

def add_run(paragraph, text, bold=False, italic=False,
            size=11, color=None, font='Calibri'):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run

def add_paragraph(doc, text='', bold=False, italic=False,
                  size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p

def add_divider(doc, color='1F3864', thickness='12'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), thickness)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def add_section_title(doc, number, title,
                      color_num='2E74B5', color_title='1F3864'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, f'{number}.  ', bold=True, size=13, color=color_num)
    add_run(p, title.upper(), bold=True, size=13, color=color_title)
    add_divider(doc, color='2E74B5', thickness='8')

def add_table_styled(doc, headers, rows,
                     hdr_bg='1F3864', hdr_fg='FFFFFF',
                     alt_bg='EBF3FB', col_widths=None):
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, hdr_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=10, color=hdr_fg)

    for r, row_data in enumerate(rows):
        bg  = alt_bg if r % 2 == 0 else 'FFFFFF'
        row = table.rows[r + 1]
        for c, text in enumerate(row_data):
            cell = row.cells[c]
            set_cell_bg(cell, bg)
            if isinstance(text, tuple):
                txt, clr, is_bold = text
                add_run(cell.paragraphs[0], str(txt), size=10, bold=is_bold, color=clr)
            else:
                add_run(cell.paragraphs[0], str(text), size=10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table

def add_callout_box(doc, title, body_lines,
                    bg='EBF3FB', border='2E74B5', title_color='1F3864'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_borders(cell, color=border)
    cell.width = Inches(6.0)

    tp = cell.add_paragraph()
    tp.paragraph_format.space_before = Pt(4)
    tp.paragraph_format.space_after  = Pt(4)
    add_run(tp, title, bold=True, size=11, color=title_color)

    for line in body_lines:
        bp = cell.add_paragraph()
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after  = Pt(3)
        bp.paragraph_format.left_indent  = Inches(0.15)
        add_run(bp, line, size=10, color='333333')

    cell.add_paragraph().paragraph_format.space_after = Pt(2)
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════
# DOCUMENTO
# ═══════════════════════════════════════════════

doc = Document()
sec = doc.sections[0]
sec.top_margin    = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)

# ── Encabezado de página ──
hdr = sec.header
hdr.is_linked_to_previous = False
ht  = hdr.add_table(1, 2, width=Inches(6.5))
ht.alignment = WD_TABLE_ALIGNMENT.CENTER
lp  = ht.cell(0, 0).paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(lp, 'Strategic Mates', bold=True, size=10, color='1F3864')
add_run(lp, '  |  SM Portal — Project Charter', size=9, color='555555')
rp  = ht.cell(0, 1).paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(rp, 'CONFIDENCIAL  –  Uso exclusivo del equipo', italic=True, size=9, color='C00000')
add_divider(hdr, color='2E74B5', thickness='6')

# ═══════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════

doc.add_paragraph()

# Logo
try:
    logo_run_p = doc.add_paragraph()
    logo_run_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run_p.paragraph_format.space_after = Pt(10)
    run = logo_run_p.add_run()
    run.add_picture(r'd:\Trabajo\Will\logo.png', height=Inches(0.65))
except Exception:
    pass

add_paragraph(doc, 'PROJECT CHARTER',
              bold=True, size=24, color='1F3864',
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

add_paragraph(doc, 'SM PORTAL',
              bold=True, size=20, color='2E74B5',
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)

add_paragraph(doc, 'Strategic Mates · Portal B2B SaaS · Fase 1',
              italic=True, size=12, color='555555',
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

add_table_styled(
    doc,
    headers=['Campo', 'Detalle'],
    rows=[
        ['Cliente',       'Strategic Mates'],
        ['Proyecto',      'SM Portal — Sistema de gestión interna B2B SaaS'],
        ['Desarrollador', 'Aquiles Dev (equipo de 1)'],
        ['Tipo',          'Migración completa + nueva plataforma'],
        ['Estado',        'Fase 1 en desarrollo activo'],
        ['Versión',       'v1.0 · Abril 2026'],
    ],
    col_widths=[1.8, 4.7]
)

# ═══════════════════════════════════════════════
# 1. DESCRIPCIÓN
# ═══════════════════════════════════════════════

add_section_title(doc, '1', 'Descripción del Proyecto')

add_callout_box(doc,
    '  ¿Qué es SM Portal?',
    [
        'Migración completa del portal operativo de Strategic Mates a una plataforma moderna.',
        'Reemplaza el sistema anterior (Firestore + HTML/PHP) con una solución robusta,',
        'segura y escalable que centraliza la gestión de franquicias, empresas SB,',
        'contratos, documentos, contabilidad y procesos internos.',
    ],
    bg='EBF3FB', border='2E74B5', title_color='1F3864'
)

# ═══════════════════════════════════════════════
# 2. ACTORES Y ROLES
# ═══════════════════════════════════════════════

add_section_title(doc, '2', 'Actores y Roles del Sistema')

add_paragraph(doc, 'Jerarquía de actores:', bold=True, size=10, color='475569', space_after=4)

for line in [
    'Strategic Mates (Superadmin — holding)',
    '    └─  SM Franchises  (ej: SM Florida, SM Texas)',
    '            └─  Small Businesses / SBs  (los clientes)',
    '                    ├─  Business Bishop  (inversor patrocinador)',
    '                    └─  Sub-Franquicias  (abiertas por el dueño del SB)',
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.2)
    indent = len(line) - len(line.lstrip())
    color = '2E74B5' if indent == 0 else ('334155' if indent < 12 else '64748B')
    add_run(p, line, size=10, color=color, bold=(indent == 0))

doc.add_paragraph()

add_table_styled(
    doc,
    headers=['Rol', 'Alcance de acceso'],
    rows=[
        ['superadmin',          'Todo el sistema'],
        ['admin_sm',            'Su franquicia SM únicamente'],
        ['sb_owner',            'Su empresa — todos los módulos'],
        ['sb_employee',         'Su empresa — módulos asignados por permisos'],
        ['bb',                  'Empresa patrocinada — Contabilidad + Contratos (solo lectura)'],
        ['sub_franchise_owner', 'Su sub-franquicia'],
        ['sub_franchise_admin', 'Su sub-franquicia (acciones de admin)'],
    ],
    col_widths=[2.2, 4.3]
)

# ═══════════════════════════════════════════════
# 3. STACK TECNOLÓGICO
# ═══════════════════════════════════════════════

add_section_title(doc, '3', 'Stack Tecnológico')

add_table_styled(
    doc,
    headers=['Capa', 'Tecnología'],
    rows=[
        ['Backend',                'Laravel 12 + PHP 8.4'],
        ['Autenticación',          'Laravel Sanctum (tokens) + Spatie Permissions (roles)'],
        ['Frontend',               'React 19 + Vite + Tailwind CSS v4'],
        ['Estado global',          'Zustand'],
        ['Enrutamiento',           'React Router v7'],
        ['Base de datos',          'PostgreSQL 16  (29 tablas: 23 de aplicación + Spatie + Sanctum)'],
        ['Caché / Colas',          'Redis 7'],
        ['Firma electrónica',      'DocuSeal (auto-hospedado)'],
        ['Editor BPMN',            'bpmn-js (bilingüe ES/EN)'],
        ['Calendario',             'FullCalendar'],
        ['Inteligencia Artificial', 'OpenAI API'],
        ['Generación de PDF',      'barryvdh/dompdf'],
        ['Infraestructura',        'Docker Compose + VPS Hostinger'],
    ],
    col_widths=[2.2, 4.3]
)

# ═══════════════════════════════════════════════
# 4. ALCANCE — FASE 1
# ═══════════════════════════════════════════════

add_section_title(doc, '4', 'Fase 1 — Alcance Actual')

add_paragraph(doc, 'Infraestructura y configuración base', bold=True, size=10, color='475569', space_after=3)
add_table_styled(
    doc,
    headers=['Componente', 'Detalle'],
    rows=[
        ['Repositorio + Docker Compose', '7 servicios: PostgreSQL, Redis, Laravel, Nginx, Queue, DocuSeal, Adminer'],
        ['Base de datos',                '29 tablas · migraciones y seeders'],
        ['Backend',                      'Laravel 12 con Sanctum, Spatie Permissions, 7 roles'],
        ['Frontend',                     'React 19 con routing, Zustand, layout completo'],
    ],
    hdr_bg='475569',
    col_widths=[2.2, 4.3]
)

add_paragraph(doc, 'Módulos del portal (15 módulos)', bold=True, size=10, color='475569', space_after=3)
add_table_styled(
    doc,
    headers=['#', 'Módulo', 'Descripción'],
    rows=[
        ['01', 'Login / Auth',            'Split layout, password toggle, tokens Sanctum'],
        ['02', 'Home / Dashboard',        'KPIs, resúmenes de actividad, accesos rápidos'],
        ['03', 'SB Applications',         'Panel de casos: scoring, análisis AI, Close Deal'],
        ['04', 'SM Franchises + Companies','CRUD de franquicias y empresas SB'],
        ['05', 'Users & Permissions',     'Gestión de usuarios y permisos por módulo'],
        ['06', 'Feed',                    'Publicaciones internas con reacciones y adjuntos'],
        ['07', 'Contracts',               'Firma digital DocuSeal · 3 firmantes por contrato'],
        ['08', 'Document Repository',     'Carga, versiones, revisión y aprobación de documentos'],
        ['09', 'Process Maps (BPMN)',     'Editor bpmn-js bilingüe ES/EN · 2 mapas por SB'],
        ['10', 'Accounting & Finance',    'Documentos contables + integración QuickBooks Online'],
        ['11', 'Inventory',               'Ítems, stock y movimientos de inventario'],
        ['12', 'Tracking',                'Proyectos, tareas y vista Gantt'],
        ['13', 'Service Catalog',         'Catálogo jerárquico (solo superadmin)'],
        ['14', 'Calendar',                'Eventos internos con FullCalendar'],
        ['15', 'User Profile',            'Datos personales, avatar, cambio de contraseña'],
    ],
    col_widths=[0.35, 2.1, 4.05]
)

# ═══════════════════════════════════════════════
# 5. ALCANCE — FASE 2
# ═══════════════════════════════════════════════

add_section_title(doc, '5', 'Fase 2 — Alcance Futuro (Post-lanzamiento)')

add_callout_box(doc,
    '  Estas funcionalidades NO están en scope de Fase 1',
    [
        '▸  Assessment público — wizard 4 etapas, 63 preguntas, 9 dimensiones, PDF de resultados (sin login)',
        '▸  Assessment 3 — formulario de auditoría',
        '▸  Formulario público BB Application',
        '▸  Flujo completo Close Deal — UI con estados + auto-creación de empresa + invitación por email',
        '▸  Landing page pública de Strategic Mates',
    ],
    bg='FFF8E1', border='C7870C', title_color='C7870C'
)

# ═══════════════════════════════════════════════
# CIERRE
# ═══════════════════════════════════════════════

add_divider(doc, color='1F3864', thickness='12')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p,
    'CONFIDENCIAL – Este documento es de uso exclusivo del equipo de proyecto y no puede '
    'ser reproducido ni distribuido sin autorización expresa.',
    italic=True, size=9, color='888888')

# ── GUARDAR ──────────────────────────────────
output_path = r'd:\Trabajo\Will\docs\PROJECT_CHARTER.docx'
doc.save(output_path)
print(f'Documento generado: {output_path}')
